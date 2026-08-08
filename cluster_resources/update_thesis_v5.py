"""
One-off script (round 3): adds a second formal hypothesis, H2, covering the
clustering analysis that was already in Section 6.5 but not tied to a stated
hypothesis. Updates 1.5 (add H2), 6.5 (state H2 is confirmed), and the
Conclusion (7.1) to reference H2 by name.

A .bak copy of the pre-edit file is written before any changes.

Run:
    python update_thesis_v5.py
"""
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v5.docx")


def find_para(doc, needle):
    return next(p for p in doc.paragraphs if needle in p.text)


def insert_body_after(anchor, text, italic=False):
    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.space_after = Pt(3)
    r = new_p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return new_p


def replace_in_run(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    raise ValueError(f"substring not found in any single run: {old!r}")


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)

    # --- 1.5 Hypothesis: add H2 right after H1 ---
    h1_para = find_para(doc, "H1: Combining visual (CLIP image)")
    insert_body_after(
        h1_para,
        "H2: Creators retrieved for the same search query are more likely to belong to the same "
        "embedding cluster than would be expected by chance, indicating that the embedding space "
        "captures meaningful, content-based groupings.",
        italic=True,
    )

    # Update the "evaluated empirically" paragraph to name H2 explicitly
    eval_para = find_para(doc, "This hypothesis is evaluated empirically in Section 6.5")
    replace_in_run(
        eval_para,
        "This hypothesis is evaluated empirically in Section 6.5",
        "H1 is evaluated empirically in Section 6.5",
    )
    replace_in_run(
        eval_para,
        "Section 6.5 additionally examines the semantic structure of the resulting creator "
        "embedding space through agglomerative clustering and silhouette analysis, to assess "
        "whether creators retrieved for the same query are grouped meaningfully.",
        "H2 is evaluated in the same section through agglomerative clustering and silhouette "
        "analysis, checking whether creators retrieved for the same query are grouped meaningfully.",
    )

    # --- 6.5: state that H2 is confirmed by the clustering results ---
    cluster_para = find_para(doc, "well above the roughly 12.5% expected by chance")
    replace_in_run(
        cluster_para,
        "This indicates that, although hypothesis H1 was not confirmed for the search-relevance "
        "metric, the underlying visual embedding space does organize creators into semantically "
        "meaningful groups that align with the tested query niches.",
        "This confirms hypothesis H2: although H1 was not confirmed for the search-relevance "
        "metric, the underlying visual embedding space does organize creators into semantically "
        "meaningful groups that align with the tested query niches.",
    )

    # --- Conclusion (7.1): reference H2 by name ---
    concl_para = find_para(doc, "A complementary clustering analysis showed")
    replace_in_run(
        concl_para,
        "A complementary clustering analysis showed that, despite modest silhouette scores, the "
        "resulting embedding space still groups creators in a way that aligns with query intent "
        "(50-80% of top-10 results sharing a common cluster per query, well above chance).",
        "A complementary clustering analysis confirmed hypothesis H2: despite modest silhouette "
        "scores, the resulting embedding space still groups creators in a way that aligns with "
        "query intent (50-80% of top-10 results sharing a common cluster per query, well above chance).",
    )

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")


if __name__ == "__main__":
    main()
