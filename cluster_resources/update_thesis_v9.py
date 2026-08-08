"""
One-off script (round 9): pre-submission review fixes.

  1. Inserts the previously-missing Fig. 6.1 image (speedup vs. CPU workers),
     which had a caption but no actual picture anywhere in the document.
  2. Discloses that the Task 1 / Task 2 evaluation corpus (Section 6.5) is
     324 of the 357 creators (33 were excluded for empty/low-quality
     transcriptions during embedding generation), and softens the "TikTok
     content lacks speech" explanation accordingly, since every excluded
     creator already lacked meaningful speech before the comparison started.
  3. Strengthens the Section 1.3 contribution bullet's hedged wording
     ("performs at least as well") to match the actual, decisive finding
     used everywhere else in the thesis ("clearly outperforms").
  4. Adds a short clarifying note reconciling the 1,000-video Section 6.3
     scalability benchmark with the 948-video Section 6.5 Whisper-fix re-run.

A .bak copy of the pre-edit file is written before any changes.

Run:
    python update_thesis_v9.py
"""
import os
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
DOCX_PATH_COPY = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft 2_main.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v9.docx")

FIG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "thesis_figures")


def find_para(doc, needle):
    return next(p for p in doc.paragraphs if needle in p.text)


def set_para_text(paragraph, new_text):
    if len(paragraph.runs) != 1:
        raise ValueError(f"expected exactly 1 run, found {len(paragraph.runs)}: {paragraph.text[:60]!r}")
    paragraph.runs[0].text = new_text


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)

    # --- 1. Insert the missing Fig. 6.1 image, right before its caption ---
    fig61_caption = find_para(doc, "Fig. 6.1. Speedup vs. number of CPU workers")
    img_p = fig61_caption.insert_paragraph_before()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(os.path.join(FIG_DIR, "fig_speedup.png"), width=Inches(5.5))
    print("Inserted missing Fig. 6.1 image")

    # --- 3. Section 1.3 contribution bullet: strengthen hedged wording ---
    p_contrib = find_para(doc, "Section 6.5 revisits this choice empirically")
    for run in p_contrib.runs:
        if "Section 6.5 revisits this choice empirically" in run.text:
            run.text = run.text.replace(
                "Section 6.5 revisits this choice empirically and finds that a visual-only "
                "weighting performs at least as well on the tested queries, so H1 (Section 1.5) "
                "is not confirmed.",
                "Section 6.5 revisits this choice empirically and finds that a visual-only "
                "weighting clearly outperforms it on the tested queries (0.48 vs. 0.28 average "
                "Precision@10), so H1 (Section 1.5) is not confirmed.",
            )
    print("Fixed Section 1.3 wording")

    # --- 2. Disclose the 324/357 evaluation-corpus gap (Section 6.5 intro) ---
    p_intro = find_para(doc, "Table 6.5 reports the resulting average Precision@10 per configuration")
    for run in p_intro.runs:
        if "for direct comparison." in run.text:
            run.text = run.text.replace(
                "for direct comparison.",
                "for direct comparison. Both this comparison and the clustering analysis below use "
                "every creator with a valid, non-empty combined embedding (n = 324 of the full "
                "357-creator corpus); the remaining 33 creators had transcriptions too short or too "
                "repetitive to pass a minimum-content quality filter applied during embedding "
                "generation and were excluded from indexing entirely, for all five configurations "
                "alike, including visual-only.",
            )
    print("Disclosed the 324/357 evaluation-corpus gap")

    # --- 2b. Soften the Whisper-fix paragraph's "lack of speech" explanation ---
    p_whisper = find_para(doc, "A natural concern is that this result merely reflects poor")
    for run in p_whisper.runs:
        if "even with clean, accurate transcriptions" in run.text:
            run.text = run.text.replace(
                "even with clean, accurate transcriptions, short-form TikTok content frequently "
                "pairs background music, jokes, or unrelated commentary with visuals unrelated to "
                "the spoken audio, so the transcript adds limited topical signal regardless of "
                "transcription accuracy.",
                "even among the 324 creators whose transcriptions were clean and did contain "
                "meaningful speech, that speech frequently describes background music, jokes, or "
                "unrelated commentary rather than the video's visual subject, so the transcript "
                "adds limited topical signal regardless of transcription accuracy.",
            )
    print("Softened the text-embedding explanation")

    # --- 2c. Task 2 paragraph: "all creator embeddings" -> explicit n = 324 ---
    p_task2_method = find_para(doc, "all creator embeddings were grouped with agglomerative")
    for run in p_task2_method.runs:
        if "all creator embeddings were grouped" in run.text:
            run.text = run.text.replace(
                "all creator embeddings were grouped with agglomerative hierarchical clustering",
                "all 324 creators with valid embeddings (see Section 6.5 introduction) were grouped "
                "with agglomerative hierarchical clustering",
            )
    print("Clarified Task 2 corpus size")

    # --- 4. Reconcile the 1,000- vs. 948-video counts ---
    for run in p_whisper.runs:
        if "the full pipeline was re-run on all 357 creators and 948 videos." in run.text:
            run.text = run.text.replace(
                "the full pipeline was re-run on all 357 creators and 948 videos.",
                "the full pipeline was re-run on all 357 creators and 948 videos (the videos present "
                "on disk at the time of this later re-run, slightly fewer than the 1,000 used for "
                "the Section 6.3 scalability benchmark).",
            )
    print("Reconciled 1,000 vs. 948 video counts")

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")
    shutil.copy(DOCX_PATH, DOCX_PATH_COPY)
    print(f"Synced copy to {DOCX_PATH_COPY}")


if __name__ == "__main__":
    main()
