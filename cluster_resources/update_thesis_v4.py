"""
One-off script (round 2): fixes internal-consistency issues identified after
reviewing the v3 draft:
  1. Conclusion (7.1) never mentioned the Section 6.5 findings -> add a paragraph.
  2. 1.3 / 3.2 reconcile the 0.85/0.15 production weights with 6.5's finding
     that visual-only scored higher (H1 not confirmed).
  3. Clarify that the 0.72 / 0.64 / 0.28 precision figures measure different
     things and are not directly comparable.
  4. Add two "Limitations" bullets to Future Work (7.2): small sample size,
     modest silhouette scores.
  5. Cite the two previously-unused references, [23] (ViT) and [24] (SBERT),
     in relevant existing sentences.

A .bak copy of the pre-edit file is written before any changes.

Run:
    python update_thesis_v4.py
"""
import os
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v4.docx")


def find_para(doc, predicate):
    return next(p for p in doc.paragraphs if predicate(p.text))


def insert_body_after(anchor, text, italic=False):
    """Inserts a justified 12pt Times New Roman body paragraph right after `anchor`."""
    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.space_after = Pt(3)
    r = new_p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return new_p


def insert_bullet_after(anchor, text):
    """Inserts a hanging-indent bullet paragraph right after `anchor`, matching
    the existing Future Work / Contributions bullet-list style."""
    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    new_p.paragraph_format.left_indent = Inches(0.39375)
    new_p.paragraph_format.first_line_indent = Inches(-0.19653)
    new_p.paragraph_format.space_after = Pt(2)
    r = new_p.add_run(f"\u2022  {text}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return new_p


def append_to_run(paragraph, addition):
    """Appends text to the paragraph's last run, inheriting its formatting."""
    last_run = paragraph.runs[-1]
    new_run = paragraph.add_run(addition)
    new_run.bold = last_run.bold
    new_run.italic = last_run.italic
    new_run.font.name = last_run.font.name
    new_run.font.size = last_run.font.size
    return new_run


def replace_in_run(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    return False


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)

    # --- Fix 1: Conclusion (7.1) should summarize the 6.5 findings ---
    concl_last = find_para(doc, lambda t: t.startswith("The Flask demo application demonstrates"))
    insert_body_after(
        concl_last,
        "Section 6.5 additionally tested hypothesis H1 (Section 1.5) by comparing five visual/text "
        "embedding weightings; contrary to expectations, the purely visual configuration outperformed "
        "all other weightings, including the deployed 85/15 configuration, so H1 was not confirmed within "
        "the scope of this evaluation. A complementary clustering analysis showed that, despite modest "
        "silhouette scores, the resulting embedding space still groups creators in a way that aligns with "
        "query intent (50-80% of top-10 results sharing a common cluster per query, well above chance). "
        "Taken together, these findings suggest that TikTok content in this domain is primarily visually "
        "driven, and that Whisper-transcription-derived text embeddings should be treated as a secondary, "
        "noisier signal rather than an equal partner to the visual signal.",
    )

    # --- Fix 2: reconcile 0.85/0.15 weights with the 6.5 finding ---
    contrib_bullet = find_para(doc, lambda t: t.startswith("\u2022  A multimodal embedding scheme"))
    append_to_run(
        contrib_bullet,
        " Section 6.5 revisits this choice empirically and finds that a visual-only weighting performs "
        "at least as well on the tested queries, so H1 (Section 1.5) is not confirmed.",
    )

    clip_weights_para = find_para(doc, lambda t: "The weights 0.85 (image) and 0.15 (text)" in t)
    append_to_run(
        clip_weights_para,
        " Section 6.5 empirically tests this choice against four alternative weightings and finds that a "
        "purely visual embedding performs comparably or better on the evaluated queries, indicating that "
        "transcription text contributes limited additional signal for this corpus.",
    )

    # Cite [23] (Vision Transformer) at its natural mention in 3.2
    clip_intro_para = find_para(doc, lambda t: t.startswith("CLIP [8] trains two encoders"))
    replace_in_run(clip_intro_para, "a Vision Transformer image encoder", "a Vision Transformer image encoder [23]")

    # --- Fix 3: clarify the 0.72 / 0.64 / 0.28 precision figures ---
    existing_es_para = find_para(doc, lambda t: t.startswith("The existing 0.85/0.15 validation result"))
    insert_body_after(
        existing_es_para,
        "To avoid ambiguity across sections: the 0.72 average precision reported in Section 6.4 measures "
        "Top-10 relevance after personalized re-ranking, a different and easier task than raw retrieval; "
        "the 0.64 and 0.28 values above both describe non-personalized retrieval on the same 85/15 "
        "configuration but differ in methodology, since 0.64 is derived from previously accumulated "
        "Elasticsearch feedback while 0.28 comes from a fresh, controlled manual relabeling of a newly "
        "built index. These three figures are therefore not directly comparable and should not be read "
        "as contradictory evidence of degraded search quality.",
    )

    # --- Fix 4: Limitations bullets in Future Work (7.2) ---
    last_bullet = find_para(doc, lambda t: t.startswith("\u2022  Explainability"))
    b1 = insert_bullet_after(
        last_bullet,
        "Larger-scale relevance evaluation: the Section 6.5 Precision@10 comparison is based on only 50 "
        "manually labeled judgments per configuration (5 queries \u00d7 10 results); a larger and more diverse "
        "query set, ideally with multiple independent annotators, would strengthen the statistical "
        "reliability of the embedding-weighting comparison.",
    )
    insert_bullet_after(
        b1,
        "Stronger clustering validation: the silhouette scores obtained in Section 6.5 (0.06-0.09) are "
        "modest by conventional standards; future work could explore domain-adapted or fine-tuned visual "
        "embeddings, alternative distance metrics, or supervised niche labels to obtain more clearly "
        "separated clusters.",
    )

    # --- Fix 5: cite [24] (Sentence-BERT) at a natural spot in 2.4 ---
    dense_retrieval_para = find_para(doc, lambda t: t.startswith("Elasticsearch 8.x [13] supports dense-vector"))
    append_to_run(
        dense_retrieval_para,
        " Reimers and Gurevych [24] similarly showed that fine-tuned Siamese sentence encoders "
        "(Sentence-BERT) produce semantically meaningful sentence embeddings suitable for efficient "
        "similarity search, an approach conceptually analogous to the CLIP text encoder used here for "
        "transcription embeddings.",
    )

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")


if __name__ == "__main__":
    main()
