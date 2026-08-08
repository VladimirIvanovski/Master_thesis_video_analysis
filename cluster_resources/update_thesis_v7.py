"""
One-off script (round 7): refreshes Section 6.5 with the re-evaluation done
after fixing two Whisper transcription bugs (repetition loops from greedy
decoding, and per-actor language-detection caching) and re-running the full
357-creator / 948-video pipeline.

Updates:
  - Table 6.5 (Precision@10 per config) with the new, content-verified numbers
  - Table 6.6 (silhouette per k) with the new k=3..15 scores
  - Figures 6.2-6.6 (replaced with regenerated PNGs from thesis_figures/)
  - Paragraph 374 (Task 1 findings) + a new paragraph describing the
    Whisper-fix robustness check
  - Paragraph 375 (comparison with historical ES numbers)
  - Paragraph 381 (Task 2 findings, new k=5 clustering)
  - Fig. 6.4 / Fig. 6.5 captions (k=8 -> k=5)
  - Paragraph 238 (wording: "comparably or better" -> "clearly outperforms")
  - Conclusion (7.1) and Future Work (7.2) numbers

A .bak copy of the pre-edit file is written before any changes.

Run:
    python update_thesis_v7.py
"""
import os
import shutil

import docx
from docx.oxml.ns import qn

DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
DOCX_PATH_COPY = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft 2_main.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v7.docx")

FIG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "thesis_figures")


def find_para(doc, needle):
    return next(p for p in doc.paragraphs if needle in p.text)


def set_para_text(paragraph, new_text):
    """Overwrites a single-run paragraph's text, keeping its formatting."""
    if len(paragraph.runs) != 1:
        raise ValueError(f"expected exactly 1 run, found {len(paragraph.runs)}: {paragraph.text[:60]!r}")
    paragraph.runs[0].text = new_text


def insert_body_after(anchor, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.space_after = Pt(3)
    r = new_p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return new_p


def replace_image(doc, paragraph, new_path):
    for run in paragraph.runs:
        for blip in run._element.findall(".//" + qn("a:blip")):
            rId = blip.get(qn("r:embed"))
            image_part = doc.part.related_parts[rId]
            with open(new_path, "rb") as f:
                image_part._blob = f.read()
            return True
    raise ValueError("no image found in paragraph")


def set_table_cell(table, row_idx, col_idx, text):
    cell = table.rows[row_idx].cells[col_idx]
    cell.paragraphs[0].runs[0].text = text


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)

    # --- Table 6.5: Precision@10 per config ---
    t5 = doc.tables[5]
    new_p10 = {"visual_only (1.00 / 0.00)": "0.52", "text_only (0.00 / 1.00)": "0.06",
               "50 / 50": "0.16", "85 / 15 (deployed)": "0.38", "15 / 85": "0.08"}
    for row in t5.rows[1:]:
        label = row.cells[0].text.strip()
        if label in new_p10:
            row.cells[1].paragraphs[0].runs[0].text = new_p10[label]
    print("Updated Table 6.5")

    # --- Table 6.6: silhouette per k ---
    t6 = doc.tables[6]
    new_sil = {3: "0.0877", 4: "0.1016", 5: "0.1067", 6: "0.1048", 7: "0.0745",
               8: "0.0819", 9: "0.0460", 10: "0.0494", 11: "0.0532", 12: "0.0603",
               13: "0.0653", 14: "0.0715", 15: "0.0753"}
    for row in t6.rows[1:]:
        k = int(row.cells[0].text.strip())
        row.cells[1].paragraphs[0].runs[0].text = new_sil[k]
    print("Updated Table 6.6")

    # --- Figures 6.2 - 6.6 ---
    fig_map = [
        (373, "fig_precision_comparison.png"),
        (380, "fig_silhouette_scores.png"),
        (383, "fig_cluster_scatter.png"),
        (385, "fig_cluster_sizes.png"),
        (387, "fig_same_cluster_pct.png"),
    ]
    for caption_idx, fname in fig_map:
        img_para = doc.paragraphs[caption_idx - 1]  # image paragraph precedes its caption
        replace_image(doc, img_para, os.path.join(FIG_DIR, fname))
    print("Replaced 5 figures")

    # --- Fig captions: k = 8 -> k = 5 ---
    for idx in (383, 385):
        p = doc.paragraphs[idx]
        set_para_text(p, p.text.replace("k = 8", "k = 5"))

    # --- Para 238: wording tweak now that the gap is decisive, not marginal ---
    p238 = doc.paragraphs[238]
    replace_target = (
        "Section 6.5 empirically tests this choice against four alternative weightings and finds "
        "that a purely visual embedding performs comparably or better on the evaluated queries, "
        "indicating that transcription text contributes limited additional signal for this corpus."
    )
    p238.runs[1].text = p238.runs[1].text.replace(
        replace_target,
        "Section 6.5 empirically tests this choice against four alternative weightings, including a "
        "re-evaluation after fixing Whisper transcription bugs, and finds that a purely visual "
        "embedding clearly outperforms every text-inclusive weighting on the evaluated queries, "
        "indicating that transcription text contributes limited additional signal for this corpus.",
    )

    # --- Para 374: Task 1 findings with new numbers ---
    set_para_text(doc.paragraphs[374], (
        "Contrary to hypothesis H1, the purely visual configuration (1.00/0.00) achieved the highest "
        "average Precision@10 (0.52) among the five tested configurations, clearly exceeding the "
        "deployed 85/15 configuration (0.38); the text-only and text-heavy configurations (0.00/1.00 "
        "and 0.15/0.85) performed substantially worse (0.06 and 0.08, respectively), and the balanced "
        "50/50 blend also underperformed the visual-only configuration (0.16). Hypothesis H1 is "
        "therefore not confirmed by this experiment: combining modalities did not outperform the best "
        "single modality. Relevance for this measurement was judged by directly inspecting a "
        "representative frame from each retrieved creator's video together with its transcription, "
        "rather than defaulting unlabeled results to \u201cnot relevant\u201d as in earlier passes, giving a "
        "more reliable precision estimate."
    ))

    # --- New paragraph: Whisper-fix robustness check ---
    insert_body_after(doc.paragraphs[375], (
        "A natural concern is that this result merely reflects poor transcription quality rather than "
        "a genuine limitation of the text modality. Manual inspection of the Whisper outputs revealed "
        "two implementation bugs: greedy decoding (beam size 1) with no repetition penalty caused the "
        "model to loop on short phrases indefinitely on ambiguous or musical audio, and the language "
        "tokenizer was cached per worker process rather than re-detected per video, causing incorrect "
        "language mixing within a batch. Both were fixed (beam size 5, a repetition penalty of 1.3 with "
        "a 3-gram repeat block, and per-item multilingual detection), and the full pipeline was re-run "
        "on all 357 creators and 948 videos. The fix visibly worked: the aggregated transcription file "
        "shrank by 74% (496 KB to 129 KB) as the repetition loops disappeared. Re-running the Task 1 "
        "evaluation on the resulting embeddings did not change the conclusion, however \u2014 the purely "
        "visual configuration remained the clear winner, this time by a wider margin (0.52 vs. 0.38 for "
        "85/15, compared with 0.32 vs. 0.28 before the fix). This indicates that the weak contribution "
        "of text embeddings is not primarily a transcription-quality artifact: even cleanly transcribed "
        "speech in short-form TikTok content frequently describes background music, jokes, or unrelated "
        "commentary rather than the video's visual subject, so it adds limited topical signal regardless "
        "of transcription accuracy."
    ))

    # --- Para (now shifted by +1): historical ES comparison ---
    p_hist = find_para(doc, "The existing 0.85/0.15 validation result (0.64) is considerably higher")
    set_para_text(p_hist, (
        "The existing 0.85/0.15 validation result (0.64) remains higher than the freshly measured value "
        "for the same configuration (0.38), though the gap narrowed considerably compared with the "
        "first pass (0.28). The remaining difference is best explained by the two evaluations testing "
        "different retrieved-creator subsets and by the historical figure being based on a small, "
        "self-selected sample of results the user had already seen and liked; it should not be read as "
        "evidence that search quality has degraded."
    ))

    # --- Para 381 (Task 2 findings): now shifted by +1 due to the inserted paragraph ---
    p381 = find_para(doc, "The silhouette score is maximized at k = 8")
    set_para_text(p381, (
        "The silhouette score is maximized at k = 5 (0.1067), yielding five clusters of size 115, 110, "
        "21, 30, and 39 creators (Fig. 6.4, Fig. 6.5). Absolute silhouette values in the 0.05-0.11 range "
        "are modest, which is expected: CLIP visual embeddings were not trained to separate TikTok "
        "content niches specifically, so cluster boundaries are soft rather than sharply delineated. "
        "Nevertheless, the clustering captures statistically real structure: for each of the five test "
        "queries, between 60% and 90% of the top-10 search results fall into a single dominant cluster "
        "(Fig. 6.6), well above the roughly 20% expected by chance with five clusters of these sizes. "
        "This confirms hypothesis H2: although H1 was not confirmed for the search-relevance metric, "
        "the underlying visual embedding space does organize creators into semantically meaningful "
        "groups that align with the tested query niches."
    ))

    # --- Conclusion (7.1) ---
    p_concl = find_para(doc, "Section 6.5 additionally tested hypothesis H1")
    set_para_text(p_concl, (
        "Section 6.5 additionally tested hypothesis H1 (Section 1.5) by comparing five visual/text "
        "embedding weightings; contrary to expectations, the purely visual configuration outperformed "
        "all other weightings, including the deployed 85/15 configuration (0.52 vs. 0.38 average "
        "Precision@10), so H1 was not confirmed within the scope of this evaluation. Fixing two Whisper "
        "transcription bugs (repetition loops and language-detection caching) and re-running the full "
        "evaluation on freshly generated embeddings did not change this conclusion, indicating the "
        "result reflects a genuine content mismatch between spoken audio and visual topic rather than a "
        "transcription-quality artifact. A complementary clustering analysis confirmed hypothesis H2: "
        "despite modest silhouette scores, the resulting embedding space still groups creators in a way "
        "that aligns with query intent (60-90% of top-10 results sharing a common cluster per query, "
        "well above chance). Taken together, these findings suggest that TikTok content in this domain "
        "is primarily visually driven, and that Whisper-transcription-derived text embeddings should be "
        "treated as a secondary, noisier signal rather than an equal partner to the visual signal."
    ))

    # --- Future Work (7.2) bullets ---
    p419 = find_para(doc, "Larger-scale relevance evaluation")
    set_para_text(p419, (
        "\u2022  Larger-scale relevance evaluation: the Section 6.5 Precision@10 comparison is based on 50 "
        "judgments per configuration (5 queries \u00d7 10 results), now judged by directly reviewing a video "
        "frame and transcription per retrieved creator rather than defaulting unlabeled rows to \u201cnot "
        "relevant\u201d; a larger and more diverse query set, ideally with multiple independent annotators, "
        "would further strengthen the statistical reliability of the embedding-weighting comparison."
    ))

    p420 = find_para(doc, "Stronger clustering validation")
    set_para_text(p420, (
        "\u2022  Stronger clustering validation: the silhouette scores obtained in Section 6.5 (0.05-0.11) "
        "are modest by conventional standards; future work could explore domain-adapted or fine-tuned "
        "visual embeddings, alternative distance metrics, or supervised niche labels to obtain more "
        "clearly separated clusters."
    ))

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")
    shutil.copy(DOCX_PATH, DOCX_PATH_COPY)
    print(f"Synced copy to {DOCX_PATH_COPY}")


if __name__ == "__main__":
    main()
