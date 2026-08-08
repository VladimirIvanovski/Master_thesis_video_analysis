"""
One-off script (round 4): adds a third formal hypothesis, H3, covering the
personalized-search evaluation that already existed in Section 6.4 (manual
Precision@10 comparison of standard vs. personalized search, per niche) but
was never framed as a tested hypothesis.

Updates 1.5 (add H3), 6.4 (state H3 is confirmed with the existing numbers),
and the Conclusion (7.1) to reference H3 by name.

A .bak copy of the pre-edit file is written before any changes.

Run:
    python update_thesis_v6.py
"""
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v6.docx")


def find_para(doc, needle):
    return next(p for p in doc.paragraphs if needle in p.text)


def insert_body_after(anchor, text, italic=False):
    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.space_after = Pt(3)
    r = new_p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return new_p


def replace_in_run(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    raise ValueError(f"substring not found in any single run: {old!r}")


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)

    # --- 1.5 Hypothesis: add H3 right after H2 ---
    h2_para = find_para(doc, "H2: Creators retrieved for the same search query")
    insert_body_after(
        h2_para,
        "H3: Personalized re-ranking based on user feedback (liked/disliked creators) improves "
        "search relevance (Precision@10) over standard, non-personalized search.",
        italic=True,
    )

    # Update the "evaluated empirically" paragraph to name H3 and where it's tested
    eval_para = find_para(doc, "H1 is evaluated empirically in Section 6.5")
    replace_in_run(
        eval_para,
        "H2 is evaluated in the same section through agglomerative clustering and silhouette "
        "analysis, checking whether creators retrieved for the same query are grouped meaningfully.",
        "H2 is evaluated in the same section through agglomerative clustering and silhouette "
        "analysis, checking whether creators retrieved for the same query are grouped meaningfully. "
        "H3 is evaluated in Section 6.4 by manually comparing standard and personalized search "
        "results across five niches.",
    )

    # --- 6.4: state that H3 is confirmed by the existing personalization results ---
    search_quality_para = find_para(doc, "Personalization consistently increased Top-10 relevance")
    replace_in_run(
        search_quality_para,
        "The largest gains appeared in Tech (+133%) and Makeup/Fashion (+100% each), while Fitness "
        "showed the smallest improvement (+50%).",
        "The largest gains appeared in Tech (+133%) and Makeup/Fashion (+100% each), while Fitness "
        "showed the smallest improvement (+50%). These results confirm hypothesis H3 (Section 1.5): "
        "personalized re-ranking improved Precision@10 in every tested niche.",
    )

    # --- Conclusion (7.1): reference H3 by name ---
    concl_para = find_para(doc, "The Flask demo application demonstrates")
    replace_in_run(
        concl_para,
        "The Flask demo application demonstrates that the resulting vector index enables natural "
        "language queries over 357 TikTok creators, with personalized reranking based on user "
        "feedback stored in Elasticsearch. Human evaluation shows consistently relevant top-10 "
        "retrievals across diverse query categories.",
        "The Flask demo application demonstrates that the resulting vector index enables natural "
        "language queries over 357 TikTok creators, with personalized reranking based on user "
        "feedback stored in Elasticsearch. Human evaluation shows consistently relevant top-10 "
        "retrievals across diverse query categories, confirming hypothesis H3: personalization "
        "improved Precision@10 in every tested niche (average 3.8/10 to 7.2/10, +92.6%).",
    )

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")


if __name__ == "__main__":
    main()
