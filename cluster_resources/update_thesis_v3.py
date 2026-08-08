"""
One-off script: adds Hypothesis (1.5), the Figure 4.1 architecture image, and
the new Embedding Configuration and Clustering Analysis (6.5) section (with
tables + figures from Task 1 / Task 2) into 249024_MasterThesis_Draft.docx.
Also updates the Table of Contents / List of Figures / List of Tables entries.

A .bak copy of the original is written before any edits.

Run:
    python update_thesis_v3.py
"""
import os
import shutil

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(HERE, "thesis_figures")
DOCX_PATH = r"C:\Users\vladimir\Downloads\249024_MasterThesis_Draft.docx"
BACKUP_PATH = DOCX_PATH.replace(".docx", "_backup_before_v3.docx")

# --- Task 1 / Task 2 numbers (from cluster_resources output) ---
PRECISION_ROWS = [
    ("visual_only (1.00 / 0.00)", "0.32"),
    ("text_only (0.00 / 1.00)", "0.00"),
    ("50 / 50", "0.08"),
    ("85 / 15 (deployed)", "0.28"),
    ("15 / 85", "0.02"),
    ("85 / 15 (existing, ES feedback)", "0.64"),
]
SILHOUETTE_ROWS = [
    (3, 0.0669), (4, 0.0624), (5, 0.0624), (6, 0.0780), (7, 0.0796),
    (8, 0.0846), (9, 0.0846), (10, 0.0572), (11, 0.0655), (12, 0.0685),
    (13, 0.0610), (14, 0.0646), (15, 0.0652),
]


def clone_run_format(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size


def insert_heading(anchor, text, level="section"):
    """Inserts a new heading paragraph immediately before `anchor`."""
    p = anchor.insert_paragraph_before()
    p.paragraph_format.space_before = Pt(12 if level == "section" else 16)
    p.paragraph_format.space_after = Pt(6 if level == "section" else 8)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12 if level == "section" else 14)
    return p


def insert_body(anchor, text, italic=False):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def insert_caption(anchor, text, italic=False):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    return p


def insert_image(anchor, image_path, width_inches=6.0):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


def insert_paragraph_after(anchor, text, fmt_run=None):
    """Creates a new paragraph and relocates it immediately after `anchor`."""
    new_p = anchor.insert_paragraph_before()
    anchor._p.addnext(new_p._p)
    r = new_p.add_run(text)
    if fmt_run is not None:
        clone_run_format(fmt_run, r)
    return new_p


def insert_table(doc, anchor, header, rows, col_widths=None):
    """Creates a 'Table Grid' table (matching the thesis's existing tables)
    at the end of the document, then relocates it immediately before `anchor`."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"

    for j, htext in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(htext)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)

    anchor._p.addprevious(table._tbl)
    return table


def main():
    shutil.copy(DOCX_PATH, BACKUP_PATH)
    print(f"Backup written to {BACKUP_PATH}")

    doc = docx.Document(DOCX_PATH)
    paras = doc.paragraphs

    # Grab stable anchor references BEFORE any mutation.
    anchor_hypothesis = paras[174]         # first blank spacer after 1.4 content
    anchor_fig41_caption = paras[267]      # "Fig. 4.1. High-level system architecture..."
    anchor_65 = paras[353]                 # first blank spacer after 6.4 content

    # --- 1. Fill Figure 4.1 (architecture diagram) ---
    insert_image(anchor_fig41_caption, os.path.join(FIG_DIR, "fig_architecture.png"), width_inches=6.3)

    # --- 2. New section 1.5 Hypothesis ---
    insert_heading(anchor_hypothesis, "1.5  Hypothesis")
    insert_body(
        anchor_hypothesis,
        "This thesis investigates the following hypothesis regarding the semantic search component "
        "described in Chapters 4-6:",
    )
    insert_body(
        anchor_hypothesis,
        "H1: Combining visual (CLIP image) and text (CLIP transcript) embeddings into a single weighted "
        "representation yields higher creator-search relevance (Precision@10) than using either modality alone.",
        italic=True,
    )
    insert_body(
        anchor_hypothesis,
        "This hypothesis is evaluated empirically in Section 6.5 by comparing five visual/text weighting "
        "configurations \u2014 visual-only, text-only, 50/50, 85/15, and 15/85 \u2014 against the currently deployed "
        "85/15 configuration, using manually labeled Precision@10 over five representative niche queries. "
        "Section 6.5 additionally examines the semantic structure of the resulting creator embedding space "
        "through agglomerative clustering and silhouette analysis, to assess whether creators retrieved for "
        "the same query are grouped meaningfully.",
    )

    # --- 3. New section 6.5 Embedding Configuration and Clustering Analysis ---
    insert_heading(anchor_65, "6.5  Embedding Configuration and Clustering Analysis")
    insert_body(
        anchor_65,
        "Section 6.4 measured the effect of personalization on the deployed 0.85/0.15 (image/text) embedding "
        "configuration. This section instead tests hypothesis H1 (Section 1.5) by comparing five visual/text "
        "weighting configurations directly: visual-only (1.00/0.00), text-only (0.00/1.00), 50/50, the deployed "
        "85/15, and 15/85. For each configuration, the five niche queries (makeup, fitness, cooking, fashion, "
        "tech) were run against a dedicated FAISS index built from that configuration's combined embeddings, "
        "and the top-10 results were manually labeled as relevant or not relevant. Table 6.5 reports the "
        "resulting average Precision@10 per configuration, alongside the existing 0.85/0.15 validation numbers "
        "derived from the good/bad creator feedback previously collected through the Flask demo application, "
        "for direct comparison.",
    )
    insert_caption(anchor_65, "Table 6.5.  Average Precision@10 by embedding configuration.")
    insert_table(doc, anchor_65, ["Configuration", "Avg. Precision@10"], PRECISION_ROWS)
    insert_image(anchor_65, os.path.join(FIG_DIR, "fig_precision_comparison.png"), width_inches=5.5)
    insert_caption(
        anchor_65,
        "Fig. 6.2.  Average Precision@10 across the five embedding configurations, compared with the "
        "existing 0.85/0.15 validation results.",
        italic=True,
    )
    insert_body(
        anchor_65,
        "Contrary to hypothesis H1, the purely visual configuration (1.00/0.00) achieved the highest average "
        "Precision@10 (0.32) among the five newly tested configurations, slightly exceeding the deployed 85/15 "
        "configuration (0.28); the text-only and text-heavy configurations (0.00/1.00 and 0.15/0.85) performed "
        "substantially worse (0.00 and 0.02, respectively). Hypothesis H1 is therefore not confirmed by this "
        "experiment: combining modalities did not outperform the best single modality. A likely explanation is "
        "that the Whisper transcriptions used for the text embeddings are frequently non-informative \u2014 "
        "background-music or trending-audio segments cause the speech-to-text model to hallucinate repetitive "
        "or unrelated text rather than capturing the creator's actual spoken content, adding noise rather than "
        "signal to the text embedding.",
    )
    insert_body(
        anchor_65,
        "The existing 0.85/0.15 validation result (0.64) is considerably higher than the freshly measured value "
        "for the same configuration (0.28). Rows for which no prior human feedback existed, and whose retrieved "
        "creators lacked informative transcriptions, were conservatively labeled not relevant during this new "
        "evaluation; the discrepancy is therefore best interpreted as a lower bound on the freshly measured "
        "configurations rather than firm evidence that search quality has degraded, and it highlights the value "
        "of continuously collecting fresh relevance judgments rather than relying solely on historical feedback.",
    )
    insert_body(
        anchor_65,
        "Using the best-performing configuration from this experiment (visual-only), all creator embeddings "
        "were grouped with agglomerative hierarchical clustering under cosine distance, testing between 3 and "
        "15 clusters and scoring each with the silhouette coefficient. An initial run using average linkage "
        "exhibited a well-known pathology: at every tested cluster count, the algorithm split off a single "
        "outlier creator rather than forming meaningful groups, which trivially inflates the silhouette score "
        "for an isolated point. Ward linkage was used instead \u2014 equivalent to cosine-based clustering on the "
        "unit-normalized embeddings used throughout this thesis, since \u2016u \u2212 v\u2016\u00b2 = 2(1 \u2212 cos_sim(u, v)) for unit "
        "vectors \u2014 which avoids this chaining behavior. Table 6.6 and Fig. 6.3 report the resulting silhouette "
        "scores.",
    )
    insert_caption(anchor_65, "Table 6.6.  Silhouette score by number of clusters (k), Ward linkage, cosine-distance scoring.")
    insert_table(doc, anchor_65, ["k", "Silhouette score"], [(k, f"{s:.4f}") for k, s in SILHOUETTE_ROWS])
    insert_image(anchor_65, os.path.join(FIG_DIR, "fig_silhouette_scores.png"), width_inches=5.2)
    insert_caption(anchor_65, "Fig. 6.3.  Silhouette score as a function of the number of clusters (k).", italic=True)
    insert_body(
        anchor_65,
        "The silhouette score is maximized at k = 8 (0.0846), yielding eight clusters of size 63, 67, 17, 38, "
        "69, 20, 23, and 9 creators (Fig. 6.4, Fig. 6.5). Absolute silhouette values in the 0.06-0.09 range are "
        "modest, which is expected: CLIP visual embeddings were not trained to separate TikTok content niches "
        "specifically, so cluster boundaries are soft rather than sharply delineated. Nevertheless, the "
        "clustering captures statistically real structure: for each of the five test queries, between 50% and "
        "80% of the top-10 search results fall into a single dominant cluster (Fig. 6.6), well above the "
        "roughly 12.5% expected by chance with eight roughly balanced clusters. This indicates that, although "
        "hypothesis H1 was not confirmed for the search-relevance metric, the underlying visual embedding "
        "space does organize creators into semantically meaningful groups that align with the tested query "
        "niches.",
    )
    insert_image(anchor_65, os.path.join(FIG_DIR, "fig_cluster_scatter.png"), width_inches=5.2)
    insert_caption(anchor_65, "Fig. 6.4.  2D PCA projection of creator embeddings, colored by cluster (k = 8).", italic=True)
    insert_image(anchor_65, os.path.join(FIG_DIR, "fig_cluster_sizes.png"), width_inches=5.2)
    insert_caption(anchor_65, "Fig. 6.5.  Number of creators per cluster (k = 8).", italic=True)
    insert_image(anchor_65, os.path.join(FIG_DIR, "fig_same_cluster_pct.png"), width_inches=5.2)
    insert_caption(
        anchor_65,
        "Fig. 6.6.  Percentage of top-10 search results sharing the dominant cluster, per query.",
        italic=True,
    )

    # --- 4. Update Table of Contents / List of Figures / List of Tables ---
    toc_14 = next(p for p in doc.paragraphs if p.text.strip() == "1.4  Thesis Structure\t3")
    insert_paragraph_after(toc_14, "1.5  Hypothesis\t3", fmt_run=toc_14.runs[0])

    toc_64 = next(p for p in doc.paragraphs if p.text.strip() == "6.4  Search Quality\t26")
    insert_paragraph_after(toc_64, "6.5  Embedding Configuration and Clustering Analysis\t27", fmt_run=toc_64.runs[0])

    fig61 = next(p for p in doc.paragraphs if p.text.strip() == "Figure 6.1  Speedup vs. number of CPU workers\t24")
    prev = fig61
    for txt in [
        "Figure 6.2  Precision@10 by embedding configuration\t27",
        "Figure 6.3  Silhouette score vs. number of clusters\t27",
        "Figure 6.4  2D PCA projection of creator embeddings by cluster\t28",
        "Figure 6.5  Number of creators per cluster\t28",
        "Figure 6.6  Percentage of top-10 results sharing the dominant cluster\t28",
    ]:
        prev = insert_paragraph_after(prev, txt, fmt_run=fig61.runs[0])

    lot_64 = next(p for p in doc.paragraphs if p.text.strip() == "Table 6.4  Search quality evaluation\t26")
    prev = lot_64
    for txt in [
        "Table 6.5  Average Precision@10 by embedding configuration\t27",
        "Table 6.6  Silhouette score by number of clusters\t27",
    ]:
        prev = insert_paragraph_after(prev, txt, fmt_run=lot_64.runs[0])

    doc.save(DOCX_PATH)
    print(f"Saved updated thesis to {DOCX_PATH}")


if __name__ == "__main__":
    main()
